"""
research_engine/cli/project_commands.py

CLI commands for the Project / Writer mode.

Commands
--------
    project new     --title TITLE --level LEVEL [--guideline FILE]
    project upload  --session SID --file FILE
    project write   --session SID --chapter N|all [--model MODEL]
    project status  --session SID
    project export  --session SID [--format docx|txt|md]
    project dataset --session SID
    project list
"""
from __future__ import annotations

import os
import sys
import json
from pathlib import Path
from datetime import datetime


# ── Colour helpers (re-exported from parent) ──────────────────
def _c(text, code):
    if not sys.stdout.isatty(): return text
    return f"\033[{code}m{text}\033[0m"

green  = lambda t: _c(t, "32")
yellow = lambda t: _c(t, "33")
red    = lambda t: _c(t, "31")
gold   = lambda t: _c(t, "33;1")
bold   = lambda t: _c(t, "1")
cyan   = lambda t: _c(t, "36")
dim    = lambda t: _c(t, "2")


# ══════════════════════════════════════════════════════════════
# Session store helpers
# ══════════════════════════════════════════════════════════════

def _sessions_dir(project_root: Path) -> Path:
    d = project_root / "sessions"
    d.mkdir(exist_ok=True)
    return d


def _session_path(project_root: Path, session_id: str) -> Path:
    return _sessions_dir(project_root) / f"{session_id}.json"


def _load_session(project_root: Path, session_id: str):
    from research_engine.writer import ProjectSession
    path = _session_path(project_root, session_id)
    if not path.exists():
        print(red(f"  Session not found: {session_id}"))
        print(dim(f"  Expected at: {path}"))
        print(dim(f"  Run: python main.py project list"))
        return None
    return ProjectSession.from_file(path)


def _save_session(session, project_root: Path):
    path = _session_path(project_root, session.session_id)
    session.save(path)
    return path


def _list_sessions(project_root: Path) -> list[Path]:
    return sorted(_sessions_dir(project_root).glob("*.json"))


# ══════════════════════════════════════════════════════════════
# project new
# ══════════════════════════════════════════════════════════════

def cmd_project_new(args, project_root: Path) -> int:
    from research_engine.writer import (
        ProjectSession, extract_text, parse_guideline,
        extract_metadata_with_ai
    )

    title     = getattr(args, "title", "") or ""
    level     = getattr(args, "level", "") or ""
    guideline = getattr(args, "guideline", None)

    print(gold("\n  📋 Research Analysis Toolkit — New Project Session"))
    print()

    session = ProjectSession.new(title=title, level=level)

    # If a guideline file was provided, load + parse it
    if guideline:
        gpath = Path(guideline)
        if not gpath.exists():
            print(red(f"  File not found: {gpath}"))
            return 1
        print(f"  📄 Parsing guideline: {gpath.name}…")
        try:
            text = extract_text(gpath)
            session.guideline_raw = text
            session.add_file(gpath.name, text)
            meta = parse_guideline(text)
            # Merge: command-line title/level override the parsed values
            if title: meta.title = title
            if level:
                from research_engine.writer import EducationLevel
                meta.level = EducationLevel.detect(level)
            session.metadata = meta
            print(green(f"  ✅ Guideline parsed ({len(text.split())} words extracted)"))
        except Exception as exc:
            print(yellow(f"  ⚠️  Could not parse guideline: {exc}"))

    # Save
    path = _save_session(session, project_root)
    print()
    print(bold("  Session created:"))
    _print_session_summary(session)
    print()
    print(dim(f"  Saved: {path}"))
    print()
    print(f"  Next step:")
    print(cyan(f"    python main.py project write --session {session.session_id} --chapter 1"))
    return 0


# ══════════════════════════════════════════════════════════════
# project upload
# ══════════════════════════════════════════════════════════════

def cmd_project_upload(args, project_root: Path) -> int:
    from research_engine.writer import extract_text, parse_guideline

    session = _load_session(project_root, args.session)
    if session is None: return 1

    fpath = Path(args.file)
    if not fpath.exists():
        print(red(f"  File not found: {fpath}"))
        return 1

    print(f"  📄 Uploading: {fpath.name}…")
    try:
        text = extract_text(fpath)
        session.add_file(fpath.name, text)
        if fpath.suffix.lower() in (".docx", ".pdf", ".txt", ".md"):
            # Also treat as guideline if no guideline yet
            if not session.guideline_raw:
                session.guideline_raw = text
                meta = parse_guideline(text)
                # Preserve existing metadata
                for attr in ("title", "level", "institution", "department",
                             "objectives", "research_questions", "hypotheses"):
                    existing = getattr(session.metadata, attr)
                    parsed   = getattr(meta, attr)
                    if not existing and parsed:
                        setattr(session.metadata, attr, parsed)
                print(green(f"  ✅ File added as project guideline"))
            else:
                print(green(f"  ✅ File added to session"))
        _save_session(session, project_root)
        print(dim(f"  Files in session: {list(session.uploaded_files.keys())}"))
    except Exception as exc:
        print(red(f"  Failed: {exc}"))
        return 1

    return 0


# ══════════════════════════════════════════════════════════════
# project write
# ══════════════════════════════════════════════════════════════

def cmd_project_write(args, project_root: Path) -> int:
    from research_engine.writer import write_chapter, CHAPTER_TITLES

    session = _load_session(project_root, args.session)
    if session is None: return 1

    model   = getattr(args, "model", "gpt-4o") or "gpt-4o"
    chapter = str(getattr(args, "chapter", "1"))

    # Determine which chapters to write
    if chapter.lower() == "all":
        to_write = session.chapters_remaining or list(range(1, 6))
    else:
        try:
            to_write = [int(chapter)]
        except ValueError:
            print(red(f"  Invalid chapter: {chapter!r} (use 1–5 or 'all')"))
            return 1

    api_key = os.environ.get("OPENAI_API_KEY", "")
    if not api_key:
        print(red("  OPENAI_API_KEY environment variable not set."))
        print(dim("  Set it with: export OPENAI_API_KEY=sk-..."))
        return 1

    for n in to_write:
        title = CHAPTER_TITLES.get(n, f"Chapter {n}")
        print(gold(f"\n  ✍️  Writing Chapter {n}: {title}…"))
        wc_target = session.metadata.word_count_for_level().get(n, 2500)
        print(dim(f"     Target: ~{wc_target:,} words | Model: {model}"))

        try:
            result = write_chapter(session, n, api_key=api_key, model=model)
            _save_session(session, project_root)
            print(green(f"  ✅ Chapter {n} written ({result.word_count:,} words)"))
            print(dim(f"     {result.model_notes}"))
            print()
            print(cyan(f"  Preview:"))
            print(dim(f"  {result.preview(300)}"))
        except Exception as exc:
            print(red(f"  ✗ Failed to write Chapter {n}: {exc}"))
            if "api_key" in str(exc).lower() or "auth" in str(exc).lower():
                return 1
            # Continue with remaining chapters

    print()
    _print_progress(session)
    return 0


# ══════════════════════════════════════════════════════════════
# project status
# ══════════════════════════════════════════════════════════════

def cmd_project_status(args, project_root: Path) -> int:
    session = _load_session(project_root, args.session)
    if session is None: return 1

    print(gold(f"\n  📊 Session: {session.session_id}"))
    print()
    _print_session_summary(session)
    print()
    _print_progress(session)
    return 0


# ══════════════════════════════════════════════════════════════
# project export
# ══════════════════════════════════════════════════════════════

def cmd_project_export(args, project_root: Path) -> int:
    from research_engine.writer import CHAPTER_TITLES

    session = _load_session(project_root, args.session)
    if session is None: return 1

    fmt      = getattr(args, "format", "docx") or "docx"
    out_dir  = project_root / "output" / f"project_{session.session_id}"
    out_dir.mkdir(parents=True, exist_ok=True)

    print(gold(f"\n  📦 Exporting project ({fmt.upper()})…"))

    if fmt == "md":
        _export_markdown(session, out_dir)
    elif fmt == "txt":
        _export_txt(session, out_dir)
    else:
        _export_docx(session, out_dir)

    print(green(f"\n  ✅ Project exported to: {out_dir}"))
    return 0


def _export_markdown(session, out_dir: Path):
    from research_engine.writer import CHAPTER_TITLES
    for n in sorted(session.chapters.keys()):
        ch   = session.chapters[n]
        path = out_dir / f"chapter_{n:02d}_{ch.title.lower().replace(' ','_')}.md"
        path.write_text(f"# Chapter {n}: {ch.title}\n\n{ch.content}", encoding="utf-8")
        print(f"  ✅ {path.name}  ({ch.word_count:,} words)")


def _export_txt(session, out_dir: Path):
    from research_engine.writer import CHAPTER_TITLES
    full_path = out_dir / f"project_{session.session_id}_full.txt"
    parts = []
    for n in sorted(session.chapters.keys()):
        ch = session.chapters[n]
        parts.append(f"\n\n{'='*70}\nCHAPTER {n}: {ch.title.upper()}\n{'='*70}\n\n{ch.content}")
    full_path.write_text("\n".join(parts), encoding="utf-8")
    total_wc = sum(ch.word_count for ch in session.chapters.values())
    print(f"  ✅ {full_path.name}  ({total_wc:,} words total)")


def _export_docx(session, out_dir: Path):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print(yellow("  python-docx not installed — exporting as Markdown instead"))
        _export_markdown(session, out_dir)
        return

    from research_engine.writer import CHAPTER_TITLES
    import re

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.54)

    # Title page
    m = session.metadata
    if m.title:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(m.title.upper())
        run.bold     = True
        run.font.size = Pt(14)
    for field_val in [m.institution, m.department, m.student_name, m.year]:
        if field_val:
            p = doc.add_paragraph(field_val)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    for n in sorted(session.chapters.keys()):
        ch    = session.chapters[n]
        lines = ch.content.split("\n")

        for line in lines:
            line_s = line.strip()
            if not line_s:
                continue
            if line_s.startswith("### "):
                h = doc.add_heading(line_s[4:], level=3)
            elif line_s.startswith("## "):
                h = doc.add_heading(line_s[3:], level=2)
            elif line_s.startswith("# "):
                h = doc.add_heading(line_s[2:], level=1)
            else:
                p = doc.add_paragraph(line_s)
                p.style.font.size = Pt(12)

        doc.add_page_break()

    path = out_dir / f"project_{session.session_id}.docx"
    doc.save(str(path))
    total_wc = sum(ch.word_count for ch in session.chapters.values())
    print(f"  ✅ {path.name}  ({total_wc:,} words total)")


# ══════════════════════════════════════════════════════════════
# project dataset
# ══════════════════════════════════════════════════════════════

def cmd_project_dataset(args, project_root: Path) -> int:
    from research_engine.writer import suggest_study_config
    import json

    session = _load_session(project_root, args.session)
    if session is None: return 1

    print(gold(f"\n  🗄️  Generating study config from project metadata…"))
    config = suggest_study_config(session)

    study_name = f"project_{session.session_id}"
    study_dir  = project_root / "studies" / study_name
    study_dir.mkdir(parents=True, exist_ok=True)

    (study_dir / "config.json").write_text(
        json.dumps(config, indent=2), encoding="utf-8")

    print(green(f"  ✅ Study config created: studies/{study_name}/config.json"))
    print()
    print(dim("  To generate the dataset, run:"))
    print(cyan(f"    python main.py run --study {study_name}"))
    print()
    print(dim("  Then add questionnaire.json and demographics.json to the study directory"))
    print(dim("  and re-run for a full analysis-ready dataset."))
    return 0


# ══════════════════════════════════════════════════════════════
# project list
# ══════════════════════════════════════════════════════════════

def cmd_project_list(project_root: Path) -> int:
    from research_engine.writer import ProjectSession

    sessions = _list_sessions(project_root)
    if not sessions:
        print(yellow("  No project sessions found."))
        print(dim(f"  Start one with: python main.py project new --title '...' --level bsc"))
        return 0

    print(gold(f"\n  📋 Project Sessions ({len(sessions)} found)\n"))
    for path in sessions:
        try:
            s = ProjectSession.from_file(path)
            done  = s.chapters_done
            title = s.metadata.title or "(untitled)"
            level = s.metadata.level.value.upper()
            print(f"  {bold(s.session_id)}  {yellow(level):<8}  {title[:50]}")
            print(dim(f"             Chapters: {done if done else 'none'} | "
                      f"Updated: {s.updated_at[:10]}"))
        except Exception:
            print(dim(f"  {path.stem}  (could not read)"))
    print()
    return 0


# ══════════════════════════════════════════════════════════════
# Print helpers
# ══════════════════════════════════════════════════════════════

def _print_session_summary(session):
    from research_engine.writer import CHAPTER_TITLES
    m = session.metadata
    rows = [
        ("Session ID",  session.session_id),
        ("Title",       m.title or "(not set)"),
        ("Level",       m.level.value.upper()),
        ("Institution", m.institution or "(not set)"),
        ("Design",      m.research_design.value.replace("_"," ").title()),
        ("Citation",    m.citation_style),
        ("Population",  (m.population or "(not set)")[:60]),
        ("Objectives",  str(len(m.objectives))),
        ("Files",       ", ".join(session.uploaded_files.keys()) or "none"),
    ]
    for label, val in rows:
        print(f"  {dim(label+':'): <22} {val}")


def _print_progress(session):
    from research_engine.writer import CHAPTER_TITLES
    print(bold("  Chapter progress:"))
    for n in range(1, 6):
        title = CHAPTER_TITLES[n]
        ch    = session.chapters.get(n)
        if ch:
            wc_str = f"  {ch.word_count:,} words"
            print(f"    {green('✅')} Ch {n}: {title}{dim(wc_str)}")
        else:
            print(f"    {dim('⬜')} Ch {n}: {dim(title)}")
